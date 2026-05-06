# -*- coding: utf-8 -*-
"""
Word 报告生成器 - 使用 python-docx 生成专业安全测试报告

Word Report Generator - Generate professional security assessment reports using python-docx
"""

import os
from typing import Dict, List, Any

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from . import (
    BaseGenerator,
    count_by_level,
    get_overall_risk_level,
    get_risk_color,
)


class WordGenerator(BaseGenerator):
    """Word 格式报告生成器"""

    # 风险等级颜色映射
    RISK_COLORS = {
        "critical": RGBColor(0xDC, 0x35, 0x45),
        "high": RGBColor(0xFD, 0x7E, 0x14),
        "medium": RGBColor(0xFF, 0xC1, 0x07),
        "low": RGBColor(0x0D, 0x6E, 0xFD),
        "info": RGBColor(0x6C, 0x75, 0x7D),
    }

    def generate(
        self,
        project_info: Dict[str, Any],
        vulnerabilities: List[Dict[str, Any]],
        output_path: str,
        lang: str = "zh",
    ) -> str:
        self.lang = lang
        doc = Document()

        self._setup_page(doc)
        self._create_cover_page(doc, project_info)
        self._create_toc_page(doc)
        self._create_chapter1_overview(doc, project_info)
        self._create_chapter2_statistics(doc, vulnerabilities)
        self._create_chapter3_vuln_details(doc, vulnerabilities)
        self._create_chapter4_risk_assessment(doc, vulnerabilities)
        self._create_chapter5_remediation_priority(doc, vulnerabilities)
        self._create_appendix(doc, vulnerabilities)

        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        doc.save(output_path)
        return output_path

    def _setup_page(self, doc: Document):
        """设置页面格式"""
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    def _set_cell_shading(self, cell, color_hex: str):
        """设置单元格背景色"""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_cell_text(self, cell, text: str, bold: bool = False, font_size: int = 10,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None):
        """设置单元格文本"""
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        run = paragraph.add_run(str(text))
        run.font.size = Pt(font_size)
        run.font.name = "Microsoft YaHei" if self.lang == "zh" else "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑" if self.lang == "zh" else "Calibri")
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = color

    def _add_heading_styled(self, doc: Document, text: str, level: int):
        """添加带样式的标题"""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = "Microsoft YaHei" if self.lang == "zh" else "Calibri"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑" if self.lang == "zh" else "Calibri")
        return heading

    def _add_paragraph_styled(self, doc: Document, text: str, font_size: int = 11,
                              bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              space_after: int = 6, color=None):
        """添加带样式的段落"""
        para = doc.add_paragraph()
        para.alignment = alignment
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = "Microsoft YaHei" if self.lang == "zh" else "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑" if self.lang == "zh" else "Calibri")
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = color
        return para

    def _create_table_with_style(self, doc: Document, rows: int, cols: int) -> Any:
        """创建带样式的表格"""
        table = doc.add_table(rows=rows, cols=cols, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        return table

    def _create_cover_page(self, doc: Document, project_info: Dict[str, Any]):
        """创建封面页"""
        # 添加空行使标题居中
        for _ in range(6):
            doc.add_paragraph()

        # 报告标题
        title = self._t("cover_title")
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        run.font.name = "Microsoft YaHei" if self.lang == "zh" else "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑" if self.lang == "zh" else "Calibri")

        # 副标题
        subtitle = self._t("cover_subtitle")
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

        # 分隔线
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("=" * 50)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        run.font.size = Pt(12)

        # 项目信息
        info_items = [
            (self._t("cover_project"), project_info.get("name", self._t("unknown"))),
            (self._t("cover_client"), project_info.get("client", self._t("unknown"))),
            (self._t("cover_tester"), project_info.get("tester", self._t("unknown"))),
            (self._t("cover_date"), project_info.get("date", "")),
            (self._t("cover_classification"), project_info.get("classification", self._t("unknown"))),
        ]

        for label, value in info_items:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(4)
            run_label = para.add_run(f"{label}: ")
            run_label.font.size = Pt(14)
            run_label.font.bold = True
            run_label.font.name = "Microsoft YaHei" if self.lang == "zh" else "Calibri"
            run_label._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑" if self.lang == "zh" else "Calibri")
            run_value = para.add_run(str(value))
            run_value.font.size = Pt(14)
            run_value.font.name = "Microsoft YaHei" if self.lang == "zh" else "Calibri"
            run_value._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑" if self.lang == "zh" else "Calibri")

        # 分页
        doc.add_page_break()

    def _create_toc_page(self, doc: Document):
        """创建目录页"""
        self._add_heading_styled(doc, self._t("toc_title"), level=1)

        toc_items = [
            self._t("ch1_title"),
            self._t("ch2_title"),
            self._t("ch3_title"),
            self._t("ch4_title"),
            self._t("ch5_title"),
            self._t("appendix_title"),
        ]

        for i, item in enumerate(toc_items, 1):
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(8)
            run = para.add_run(f"{i}.  {item}")
            run.font.size = Pt(12)
            run.font.name = "Microsoft YaHei" if self.lang == "zh" else "Calibri"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑" if self.lang == "zh" else "Calibri")

        doc.add_page_break()

    def _create_chapter1_overview(self, doc: Document, project_info: Dict[str, Any]):
        """第一章：项目概述"""
        self._add_heading_styled(doc, self._t("ch1_title"), level=1)

        # 测试目标
        self._add_heading_styled(doc, self._t("ch1_test_objective"), level=2)
        self._add_paragraph_styled(doc, self._t("ch1_test_objective_desc"))

        # 测试范围
        self._add_heading_styled(doc, self._t("ch1_test_scope"), level=2)
        scope_text = project_info.get("scope", self._t("ch1_test_scope_desc"))
        if isinstance(scope_text, list):
            for item in scope_text:
                self._add_paragraph_styled(doc, f"  - {item}", font_size=11)
        else:
            self._add_paragraph_styled(doc, str(scope_text))

        # 测试方法
        self._add_heading_styled(doc, self._t("ch1_test_method"), level=2)
        method_text = project_info.get("method", self._t("ch1_test_method_desc"))
        if isinstance(method_text, list):
            for item in method_text:
                self._add_paragraph_styled(doc, f"  - {item}", font_size=11)
        else:
            self._add_paragraph_styled(doc, str(method_text))

        # 测试工具
        self._add_heading_styled(doc, self._t("ch1_test_tools"), level=2)
        tools_text = project_info.get("tools", self._t("ch1_test_tools_desc"))
        if isinstance(tools_text, list):
            for item in tools_text:
                self._add_paragraph_styled(doc, f"  - {item}", font_size=11)
        else:
            self._add_paragraph_styled(doc, str(tools_text))

        doc.add_page_break()

    def _create_chapter2_statistics(self, doc: Document, vulnerabilities: List[Dict[str, Any]]):
        """第二章：风险统计概览"""
        self._add_heading_styled(doc, self._t("ch2_title"), level=1)
        self._add_paragraph_styled(doc, self._t("ch2_desc"))

        counts = count_by_level(vulnerabilities)
        total = sum(counts.values())

        # 统计表
        table = self._create_table_with_style(doc, rows=7, cols=3)

        # 表头
        headers = [self._t("ch2_level"), self._t("ch2_count"), self._t("ch2_percentage")]
        for j, header in enumerate(headers):
            self._set_cell_text(table.rows[0].cells[j], header, bold=True, font_size=10,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(table.rows[0].cells[j], "1A237E")
            for paragraph in table.rows[0].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # 数据行
        levels = [
            ("critical", self._t("ch2_critical")),
            ("high", self._t("ch2_high")),
            ("medium", self._t("ch2_medium")),
            ("low", self._t("ch2_low")),
            ("info", self._t("ch2_info")),
        ]

        for i, (level_key, level_name) in enumerate(levels):
            row = table.rows[i + 1]
            count = counts[level_key]
            percentage = f"{count / total * 100:.1f}%" if total > 0 else "0.0%"
            self._set_cell_text(row.cells[0], level_name, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_text(row.cells[1], str(count), font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_text(row.cells[2], percentage, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            # 设置风险等级颜色
            color = self.RISK_COLORS.get(level_key, RGBColor(0x6C, 0x75, 0x7D))
            self._set_cell_shading(row.cells[0], get_risk_color(level_key))
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True

        # 合计行
        total_row = table.rows[6]
        self._set_cell_text(total_row.cells[0], self._t("ch2_total"), bold=True, font_size=10,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_text(total_row.cells[1], str(total), bold=True, font_size=10,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_text(total_row.cells[2], "100.0%", bold=True, font_size=10,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(total_row.cells[0], "E8EAF6")
        self._set_cell_shading(total_row.cells[1], "E8EAF6")
        self._set_cell_shading(total_row.cells[2], "E8EAF6")

        doc.add_page_break()

    def _create_chapter3_vuln_details(self, doc: Document, vulnerabilities: List[Dict[str, Any]]):
        """第三章：漏洞详情"""
        self._add_heading_styled(doc, self._t("ch3_title"), level=1)

        if not vulnerabilities:
            self._add_paragraph_styled(doc, self._t("html_no_vulns"), font_size=12)
            return

        for idx, vuln in enumerate(vulnerabilities, 1):
            # 漏洞标题
            vuln_name = vuln.get("name", self._t("unknown"))
            risk_level = str(vuln.get("risk_level", "info")).lower()
            risk_text = self._risk_text(risk_level)

            heading_text = f"{idx}. {vuln_name} [{risk_text}]"
            self._add_heading_styled(doc, heading_text, level=2)

            # 漏洞信息表
            table = self._create_table_with_style(doc, rows=7, cols=2)

            info_items = [
                (self._t("ch3_vuln_id"), str(idx).zfill(4)),
                (self._t("ch3_cve_id"), vuln.get("cve_id", self._t("na"))),
                (self._t("ch3_vuln_name"), vuln_name),
                (self._t("ch3_risk_level"), risk_text),
                (self._t("ch3_target"), vuln.get("target", self._t("na"))),
                (self._t("ch3_port"), str(vuln.get("port", self._t("na")))),
                (self._t("ch3_protocol"), vuln.get("protocol", self._t("na"))),
            ]

            for i, (label, value) in enumerate(info_items):
                self._set_cell_text(table.rows[i].cells[0], label, bold=True, font_size=10)
                self._set_cell_shading(table.rows[i].cells[0], "E8EAF6")
                self._set_cell_text(table.rows[i].cells[1], value, font_size=10)

            # 设置风险等级颜色
            risk_color = self.RISK_COLORS.get(risk_level, RGBColor(0x6C, 0x75, 0x7D))
            for paragraph in table.rows[3].cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = risk_color
                    run.bold = True

            doc.add_paragraph()  # 空行

            # 描述
            self._add_paragraph_styled(doc, self._t("ch3_description"), font_size=11, bold=True)
            self._add_paragraph_styled(doc, vuln.get("description", self._t("na")), font_size=10)

            # 影响分析
            self._add_paragraph_styled(doc, self._t("ch3_impact"), font_size=11, bold=True)
            self._add_paragraph_styled(doc, vuln.get("impact", self._t("na")), font_size=10)

            # 复现步骤
            reproduce_steps = vuln.get("reproduce_steps", "")
            if reproduce_steps:
                self._add_paragraph_styled(doc, self._t("ch3_reproduce"), font_size=11, bold=True)
                if isinstance(reproduce_steps, list):
                    for step_i, step in enumerate(reproduce_steps, 1):
                        self._add_paragraph_styled(doc, f"  {step_i}. {step}", font_size=10)
                else:
                    self._add_paragraph_styled(doc, f"  {reproduce_steps}", font_size=10)

            # 修复建议
            self._add_paragraph_styled(doc, self._t("ch3_remediation"), font_size=11, bold=True)
            self._add_paragraph_styled(doc, vuln.get("remediation", self._t("na")), font_size=10)

            # 来源
            source = vuln.get("source", self._t("na"))
            self._add_paragraph_styled(doc, f"{self._t('ch3_source')}: {source}",
                                       font_size=9, color=RGBColor(0x6C, 0x75, 0x7D))

            doc.add_paragraph()  # 漏洞间空行

        doc.add_page_break()

    def _create_chapter4_risk_assessment(self, doc: Document, vulnerabilities: List[Dict[str, Any]]):
        """第四章：综合风险评估"""
        self._add_heading_styled(doc, self._t("ch4_title"), level=1)

        counts = count_by_level(vulnerabilities)
        overall = get_overall_risk_level(counts)
        overall_text = self._risk_text(overall)

        # 整体风险等级
        self._add_heading_styled(doc, self._t("ch4_overall_risk"), level=2)
        risk_color = self.RISK_COLORS.get(overall, RGBColor(0x6C, 0x75, 0x7D))
        para = self._add_paragraph_styled(doc, overall_text, font_size=18, bold=True,
                                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for run in para.runs:
            run.font.color.rgb = risk_color

        # 风险概述
        self._add_heading_styled(doc, self._t("ch4_risk_summary"), level=2)
        self._add_paragraph_styled(doc, self._t("ch4_risk_summary_desc"))

        # 各等级风险描述
        level_descriptions = [
            ("critical", self._t("ch4_critical_desc")),
            ("high", self._t("ch4_high_desc")),
            ("medium", self._t("ch4_medium_desc")),
            ("low", self._t("ch4_low_desc")),
            ("info", self._t("ch4_info_desc")),
        ]

        for level_key, desc in level_descriptions:
            if counts[level_key] > 0:
                level_name = self._risk_text(level_key)
                count = counts[level_key]
                self._add_paragraph_styled(
                    doc, f"[{level_name}] ({count})", font_size=11, bold=True,
                    color=self.RISK_COLORS.get(level_key, RGBColor(0, 0, 0))
                )
                self._add_paragraph_styled(doc, desc, font_size=10)

        # 综合建议
        self._add_heading_styled(doc, self._t("ch4_recommendation"), level=2)
        self._add_paragraph_styled(doc, self._t("ch4_recommendation_desc"))

        doc.add_page_break()

    def _create_chapter5_remediation_priority(self, doc: Document, vulnerabilities: List[Dict[str, Any]]):
        """第五章：修复优先级建议"""
        self._add_heading_styled(doc, self._t("ch5_title"), level=1)
        self._add_paragraph_styled(doc, self._t("ch5_desc"))

        counts = count_by_level(vulnerabilities)

        # 优先级表
        table = self._create_table_with_style(doc, rows=5, cols=4)

        # 表头
        headers = [self._t("ch5_priority"), self._t("ch5_level"), self._t("ch5_count"), self._t("ch5_suggestion")]
        for j, header in enumerate(headers):
            self._set_cell_text(table.rows[0].cells[j], header, bold=True, font_size=10,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(table.rows[0].cells[j], "1A237E")
            for paragraph in table.rows[0].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # 优先级数据
        priorities = [
            (self._t("ch5_p1"), self._t("ch2_critical"), counts["critical"], self._t("ch5_p1_desc"), "critical"),
            (self._t("ch5_p2"), self._t("ch2_high"), counts["high"], self._t("ch5_p2_desc"), "high"),
            (self._t("ch5_p3"), self._t("ch2_medium"), counts["medium"], self._t("ch5_p3_desc"), "medium"),
            (self._t("ch5_p4"), f"{self._t('ch2_low')}/{self._t('ch2_info')}",
             counts["low"] + counts["info"], self._t("ch5_p4_desc"), "low"),
        ]

        for i, (priority, level, count, suggestion, color_key) in enumerate(priorities):
            row = table.rows[i + 1]
            self._set_cell_text(row.cells[0], priority, bold=True, font_size=10,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_text(row.cells[1], level, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_text(row.cells[2], str(count), font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_text(row.cells[3], suggestion, font_size=9)
            # 颜色标记
            bg_color = get_risk_color(color_key)
            self._set_cell_shading(row.cells[1], bg_color)
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True

        doc.add_page_break()

    def _create_appendix(self, doc: Document, vulnerabilities: List[Dict[str, Any]]):
        """附录：漏洞统计汇总表"""
        self._add_heading_styled(doc, self._t("appendix_title"), level=1)

        if not vulnerabilities:
            self._add_paragraph_styled(doc, self._t("html_no_vulns"), font_size=12)
            return

        # 创建汇总表
        table = self._create_table_with_style(doc, rows=len(vulnerabilities) + 1, cols=8)

        # 表头
        headers = [
            self._t("appendix_index"),
            self._t("appendix_cve"),
            self._t("appendix_name"),
            self._t("appendix_level"),
            self._t("appendix_target"),
            self._t("appendix_port"),
            self._t("appendix_protocol"),
            self._t("appendix_source"),
        ]
        for j, header in enumerate(headers):
            self._set_cell_text(table.rows[0].cells[j], header, bold=True, font_size=9,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(table.rows[0].cells[j], "1A237E")
            for paragraph in table.rows[0].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # 数据行
        for i, vuln in enumerate(vulnerabilities):
            row = table.rows[i + 1]
            risk_level = str(vuln.get("risk_level", "info")).lower()

            values = [
                str(i + 1),
                vuln.get("cve_id", self._t("na")),
                vuln.get("name", self._t("unknown")),
                self._risk_text(risk_level),
                vuln.get("target", self._t("na")),
                str(vuln.get("port", self._t("na"))),
                vuln.get("protocol", self._t("na")),
                vuln.get("source", self._t("na")),
            ]

            for j, value in enumerate(values):
                self._set_cell_text(row.cells[j], value, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # 风险等级颜色
            self._set_cell_shading(row.cells[3], get_risk_color(risk_level))
            for paragraph in row.cells[3].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
